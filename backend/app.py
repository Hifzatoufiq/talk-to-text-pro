import os
import tempfile
from datetime import datetime, timezone
from functools import wraps
from io import BytesIO

from flask import (
    Flask, request, render_template, redirect, url_for,
    flash, jsonify, send_file, session
)
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash

from pymongo import MongoClient
from bson import ObjectId

# Core utils
from utils.audio_processing import extract_audio_from_file, transcribe_audio, summarize_text

# OpenAI (for Translator)
from openai import OpenAI

# OCR / PDF / DOCX utils
import fitz  # PyMuPDF
import pytesseract
from PIL import Image, ImageOps  # ImageOps for simple preprocessing
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

# OPTIONAL (Windows): set Tesseract installed path; comment out on mac/linux
import platform
if platform.system() == "Windows":
    default_tess = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    if os.path.exists(default_tess):
        pytesseract.pytesseract.tesseract_cmd = default_tess

# ---------------- Paths ----------------
BASE_DIR = os.path.abspath(os.path.dirname(__file__))
TEMPLATES_DIR = os.path.join(BASE_DIR, "templates")
STATIC_DIR = os.path.join(BASE_DIR, "static")
UPLOAD_FOLDER = os.path.join(STATIC_DIR, "uploads")
ALLOWED_EXTENSIONS = {'mp3', 'wav', 'm4a', 'mp4', 'opus', 'ogg', 'webm', 'mov', 'mkv', 'avi', 'aac'}

app = Flask(__name__, template_folder=TEMPLATES_DIR, static_folder=STATIC_DIR)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["SECRET_KEY"] = "devsecret"
os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)

# ---------------- Mongo ----------------
MONGO_URI = os.getenv("MONGO_URI", "mongodb://localhost:27017")
client = MongoClient(MONGO_URI)
db = client["talktodb"]
users_collection         = db["users"]
notes_collection         = db["notes"]
translations_collection  = db["translations"]
conversions_collection   = db["conversions"]

# ---------------- OpenAI ----------------
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY") or os.getenv("OPENAI_API_KEY_PUBLIC") or ""
if not OPENAI_API_KEY:
    raise RuntimeError("OPENAI_API_KEY is not set. Please set it in the environment or .env.")
oai = OpenAI(api_key=OPENAI_API_KEY)

# ---------------- Helpers ----------------
def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

def login_required(view_fn):
    @wraps(view_fn)
    def wrapper(*args, **kwargs):
        if not session.get("user_id"):
            flash("Please log in to use Meeting Notes.", "warning")
            return redirect(url_for("login", next=request.path))
        return view_fn(*args, **kwargs)
    return wrapper

def _pick_ocr_langs() -> str:
    """
    Decide which OCR languages to use:
    - If env TESS_LANG is set, use that.
    - Else pick from installed langs (eng/urd/hin/ara if available).
    - Always fall back to 'eng' to avoid Tesseract errors.
    """
    env_langs = (os.getenv("TESS_LANG") or "").strip()
    if env_langs:
        return env_langs

    try:
        available = set(pytesseract.get_languages(config=""))
    except Exception:
        available = {"eng"}

    # prefer these if installed (order matters)
    preferred = ["eng", "urd", "hin", "ara"]
    chosen = [l for l in preferred if l in available]
    if not chosen:
        chosen = ["eng"]
    # make a unique "+"-joined list preserving order
    seen = {}
    return "+".join([seen.setdefault(x, x) for x in chosen if x not in seen])

def extract_pdf_text(file_path: str) -> str:
    """
    Robust PDF text extraction with high-DPI OCR fallback.
    1) Try the embedded text layer.
    2) If empty (scanned/screenshot page), render at ~250 DPI and OCR.
    3) Light preprocessing (grayscale + autocontrast) for better OCR.
    """
    text_chunks = []
    doc = fitz.open(file_path)
    ocr_langs = _pick_ocr_langs()  # e.g., "eng+urd+hin"

    for page in doc:
        # Try embedded text first
        page_text = page.get_text("text").strip()
        if page_text:
            text_chunks.append(page_text)
            continue

        # OCR fallback for image-only pages
        try:
            zoom = 3.5  # 72 * 3.5 ≈ 252 DPI
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat, alpha=False)

            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            gray = ImageOps.grayscale(img)
            gray = ImageOps.autocontrast(gray)

            # PSM 6: Assume a single uniform block of text
            ocr_text = pytesseract.image_to_string(gray, lang=ocr_langs, config="--psm 6").strip()
        except Exception:
            ocr_text = ""

        text_chunks.append(ocr_text)

    doc.close()
    # Join non-empty chunks only
    return "\n\n".join([t for t in text_chunks if t]).strip()

def translate_with_openai(text: str, target_language: str) -> str:
    """
    Simple translator via OpenAI chat.
    """
    if not text.strip():
        return ""
    prompt = f"Translate the following text to {target_language}. Keep meaning, tone, and formatting.\n\n{text}"
    resp = oai.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are a world-class translator."},
            {"role": "user", "content": prompt},
        ],
        temperature=0.2,
    )
    return (resp.choices[0].message.content if resp.choices else "") or ""

# ---------------- Public pages ----------------
@app.route("/")
def index():
    return render_template("index.html")

@app.route("/about")
def about():
    return render_template("about.html")

# ---------------- Auth ----------------
@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        name = (request.form.get("name") or "").strip()
        email = (request.form.get("email") or "").strip().lower()
        password = request.form.get("password") or ""

        if not name or not email or not password:
            flash("All fields are required.", "error")
            return redirect(url_for("register"))

        if users_collection.find_one({"email": email}):
            flash("Email already registered. Please login.", "warning")
            return redirect(url_for("login"))

        users_collection.insert_one({
            "name": name,
            "email": email,
            "password_hash": generate_password_hash(password),
            "created_at": datetime.now(timezone.utc),
        })

        user = users_collection.find_one({"email": email})
        session["user_id"] = str(user["_id"])
        session["user_name"] = user.get("name")
        session["user_email"] = user.get("email")

        flash("Welcome! Account created.", "success")
        return redirect(url_for("index"))

    return render_template("register.html")

@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        email = (request.form.get("email") or "").strip().lower()
        password = request.form.get("password") or ""
        next_url = request.args.get("next") or url_for("index")

        user = users_collection.find_one({"email": email})
        if not user or not check_password_hash(user["password_hash"], password):
            flash("Invalid email or password.", "error")
            return redirect(url_for("login", next=next_url))

        session["user_id"] = str(user["_id"])
        session["user_name"] = user.get("name")
        session["user_email"] = user.get("email")
        flash("Logged in.", "success")
        return redirect(next_url)

    return render_template("login.html")

@app.route("/logout")
def logout():
    session.clear()
    flash("Logged out.", "success")
    return redirect(url_for("index"))

# ---------------- Meeting Notes ----------------
@app.route("/upload", methods=["POST"])
@login_required
def upload():
    """
    Upload audio/video and generate notes.
    If user selected Auto, summarize in English.
    If user selected a specific language, summarize in that language.
    """
    if "audio" not in request.files:
        flash("No audio file part", "error")
        return redirect(url_for("index"))

    file = request.files["audio"]
    language = request.form.get("language", "auto")

    if file.filename == "":
        flash("No selected file", "error")
        return redirect(url_for("index"))

    if not (file and allowed_file(file.filename)):
        flash("File type not allowed", "error")
        return redirect(url_for("index"))

    filename = secure_filename(file.filename)
    save_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
    os.makedirs(os.path.dirname(save_path), exist_ok=True)
    file.save(save_path)

    try:
        audio_path = extract_audio_from_file(save_path)

        # Transcribe (auto-detect if language == auto)
        transcript = transcribe_audio(audio_path, language=language)
        if not transcript:
            transcript = "⚠️ Transcription failed or empty."

        # Summarize: Auto -> English, else chosen language
        summary_lang = "en" if language.lower() == "auto" else language
        summary = summarize_text(transcript, target_language=summary_lang)
        if not summary:
            summary = {
                "executive_summary": "⚠️ Failed to generate summary",
                "key_points": [],
                "action_items": [],
                "decisions": [],
                "sentiment": "Unknown"
            }

        note_doc = {
            "filename": filename,
            "converted_wav": os.path.basename(audio_path),
            "language": language,
            "transcript": transcript,
            "summary": summary,
            "created_at": datetime.now(timezone.utc),
            "owner_id": session.get("user_id"),
        }
        inserted = notes_collection.insert_one(note_doc)
        note_id = str(inserted.inserted_id)

        note_doc["_id"] = note_id
        return render_template("notes.html", result=note_doc)
    except Exception as e:
        flash(f"Processing failed: {e}", "error")
        return redirect(url_for("index"))

@app.route("/record", methods=["POST"])
@login_required
def record():
    """
    Record from mic/tab and generate notes.
    Respect dropdown: Auto => English notes, others => selected language.
    """
    if "audio" not in request.files:
        return jsonify({"error": "No audio data received"}), 400

    language = request.form.get("language", "auto")
    file = request.files["audio"]

    suffix = os.path.splitext(secure_filename(file.filename))[1] or ".webm"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        file.save(tmp.name)
        temp_path = tmp.name

    try:
        audio_path = extract_audio_from_file(temp_path)

        transcript = transcribe_audio(audio_path, language=language)
        if not transcript:
            transcript = "⚠️ Transcription failed or empty."

        summary_lang = "en" if language.lower() == "auto" else language
        summary = summarize_text(transcript, target_language=summary_lang)
        if not summary:
            summary = {
                "executive_summary": "⚠️ Failed to generate summary",
                "key_points": [],
                "action_items": [],
                "decisions": [],
                "sentiment": "Unknown"
            }

        note_doc = {
            "filename": os.path.basename(temp_path),
            "converted_wav": os.path.basename(audio_path),
            "language": language,
            "transcript": transcript,
            "summary": summary,
            "created_at": datetime.now(timezone.utc),
            "owner_id": session.get("user_id"),
        }
        inserted = notes_collection.insert_one(note_doc)
        note_id = str(inserted.inserted_id)

        redirect_url = url_for("view_note", note_id=note_id)
        return jsonify({"redirect_url": redirect_url})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/notes/<note_id>")
@login_required
def view_note(note_id):
    try:
        doc = notes_collection.find_one({"_id": ObjectId(note_id)})
        if not doc:
            return "Note not found", 404
        if doc.get("owner_id") and doc["owner_id"] != session.get("user_id"):
            return "Forbidden", 403

        doc["_id"] = str(doc["_id"])
        return render_template("notes.html", result=doc)
    except Exception as e:
        return f"Error: {e}", 500

# ---------------- Download Notes (TXT / DOCX) ----------------
def _build_docx_from_note(note) -> BytesIO:
    """
    Create a DOCX in-memory from a note document and return a BytesIO stream.
    """
    bio = BytesIO()
    d = Document()

    d.add_heading("Meeting Notes", 0)
    d.add_paragraph(f"File: {note.get('filename','')}")
    d.add_paragraph(f"Language: {note.get('language','')}")
    d.add_paragraph(f"Created: {note.get('created_at','')}")
    d.add_paragraph("")

    s = note.get("summary", {}) or {}

    d.add_heading("Executive Summary", level=1)
    d.add_paragraph(s.get("executive_summary","") or "")

    d.add_heading("Key Points", level=1)
    for p in (s.get("key_points") or []):
        d.add_paragraph(p, style="List Bullet")

    d.add_heading("Action Items", level=1)
    for a in (s.get("action_items") or []):
        if isinstance(a, dict):
            line = f"{a.get('task','')}"
            if a.get("owner"): line += f" — {a['owner']}"
            if a.get("due"):   line += f" (Due: {a['due']})"
            d.add_paragraph(line, style="List Bullet")
        else:
            d.add_paragraph(str(a), style="List Bullet")

    d.add_heading("Decisions", level=1)
    for dec in (s.get("decisions") or []):
        d.add_paragraph(dec, style="List Bullet")

    d.add_heading("Sentiment", level=1)
    d.add_paragraph(s.get("sentiment","Unknown"))

    d.add_heading("Transcript", level=1)
    d.add_paragraph(note.get("transcript","") or "")

    d.save(bio)
    bio.seek(0)
    return bio

@app.route("/notes/<note_id>/download")
@login_required
def download_note(note_id):
    """
    Download the note as TXT (default) or DOCX (?format=docx).
    """
    fmt = (request.args.get("format") or "txt").lower()
    doc = notes_collection.find_one({"_id": ObjectId(note_id)})
    if not doc:
        return "Note not found", 404
    if doc.get("owner_id") and doc["owner_id"] != session.get("user_id"):
        return "Forbidden", 403

    if fmt == "docx":
        bio = _build_docx_from_note({**doc, "_id": str(doc["_id"])})
        return send_file(
            bio,
            as_attachment=True,
            download_name=f"notes_{doc['_id']}.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # TXT fallback
    s = doc.get("summary", {}) or {}
    lines = []
    lines.append("MEETING NOTES")
    lines.append(f"File: {doc.get('filename','')}")
    lines.append(f"Language: {doc.get('language','')}")
    lines.append(f"Created: {doc.get('created_at','')}")
    lines.append("")
    lines.append("Executive Summary")
    lines.append(s.get("executive_summary","") or "")
    lines.append("")
    lines.append("Key Points")
    for p in (s.get("key_points") or []): lines.append(f"- {p}")
    lines.append("")
    lines.append("Action Items")
    for a in (s.get("action_items") or []):
        if isinstance(a, dict):
            line = f"- {a.get('task','')}"
            if a.get("owner"): line += f" — {a['owner']}"
            if a.get("due"):   line += f" (Due: {a['due']})"
            lines.append(line)
        else:
            lines.append(f"- {a}")
    lines.append("")
    lines.append("Decisions")
    for d in (s.get("decisions") or []): lines.append(f"- {d}")
    lines.append("")
    lines.append(f"Sentiment: {s.get('sentiment','Unknown')}")
    lines.append("")
    lines.append("Transcript")
    lines.append(doc.get("transcript","") or "")

    txt_bytes = BytesIO("\n".join(lines).encode("utf-8"))
    txt_bytes.seek(0)
    return send_file(
        txt_bytes,
        as_attachment=True,
        download_name=f"notes_{doc['_id']}.txt",
        mimetype="text/plain; charset=utf-8"
    )

# ---------------- Translator & Converters (Public) ----------------
@app.route("/translator", methods=["GET", "POST"])
def translator():
    """
    Public translator page:
    - Translate raw text
    - Translate a PDF (OCR-aware)
    """
    user_input = ""
    translated_text = ""
    target_language = request.form.get("target_lang", "Urdu")

    if request.method == "POST":
        # Text translation
        if "message" in request.form and request.form["message"].strip():
            user_input = request.form["message"]
            try:
                translated_text = translate_with_openai(user_input, target_language)
            except Exception as e:
                flash(f"Translation failed: {e}", "error")
                translated_text = ""

            translations_collection.insert_one({
                "type": "text",
                "source_text": user_input,
                "translated_text": translated_text,
                "target_language": target_language,
                "created_at": datetime.now(timezone.utc),
                "owner_id": session.get("user_id")
            })

        # PDF translation
        elif "pdf_file" in request.files:
            pdf_file = request.files["pdf_file"]
            if pdf_file and pdf_file.filename != "":
                in_name = secure_filename(pdf_file.filename)
                in_path = os.path.join(app.config["UPLOAD_FOLDER"], in_name)
                os.makedirs(os.path.dirname(in_path), exist_ok=True)
                pdf_file.save(in_path)
                try:
                    pdf_text = extract_pdf_text(in_path)
                except Exception as e:
                    flash(f"Failed to read PDF: {e}", "error")
                    pdf_text = ""

                if not pdf_text.strip():
                    translated_text = "Could not extract text from the PDF."
                    user_input = ""
                else:
                    user_input = pdf_text
                    try:
                        translated_text = translate_with_openai(pdf_text[:15000], target_language)
                    except Exception as e:
                        flash(f"Translation failed: {e}", "error")
                        translated_text = ""

                translations_collection.insert_one({
                    "type": "pdf",
                    "filename": in_name,
                    "source_text": user_input,
                    "translated_text": translated_text,
                    "target_language": target_language,
                    "created_at": datetime.now(timezone.utc),
                    "owner_id": session.get("user_id")
                })

    return render_template(
        "translator.html",
        user_input=user_input,
        translated_text=translated_text,
        target_language=target_language,
    )

@app.route("/pdf-to-docx", methods=["POST"])
def pdf_to_docx():
    pdf_file = request.files.get("pdf_file")
    if not pdf_file or pdf_file.filename == "":
        flash("Please select a PDF file.", "error")
        return redirect(url_for("translator"))

    in_name = secure_filename(pdf_file.filename)
    in_path = os.path.join(app.config["UPLOAD_FOLDER"], in_name)
    os.makedirs(os.path.dirname(in_path), exist_ok=True)
    pdf_file.save(in_path)

    try:
        pdf_text = extract_pdf_text(in_path)
        out_path = os.path.splitext(in_path)[0] + ".docx"
        os.makedirs(os.path.dirname(out_path), exist_ok=True)

        doc = Document()
        for line in pdf_text.splitlines():
            doc.add_paragraph(line)
        doc.save(out_path)

        if not os.path.isfile(out_path):
            flash("Failed to create DOCX file.", "error")
            return redirect(url_for("translator"))

        conversions_collection.insert_one({
            "type": "pdf-to-docx",
            "src_filename": in_name,
            "out_filename": os.path.basename(out_path),
            "created_at": datetime.now(timezone.utc),
            "owner_id": session.get("user_id")
        })

        return send_file(out_path, as_attachment=True, download_name=os.path.basename(out_path))
    except Exception as e:
        flash(f"PDF to DOCX failed: {e}", "error")
        return redirect(url_for("translator"))

@app.route("/docx-to-pdf", methods=["POST"])
def docx_to_pdf():
    docx_file = request.files.get("docx_file")
    if not docx_file or docx_file.filename == "":
        flash("Please select a DOCX file.", "error")
        return redirect(url_for("translator"))

    in_name = secure_filename(docx_file.filename)
    in_path = os.path.join(app.config["UPLOAD_FOLDER"], in_name)
    os.makedirs(os.path.dirname(in_path), exist_ok=True)
    docx_file.save(in_path)

    try:
        doc = Document(in_path)
        out_path = os.path.splitext(in_path)[0] + ".pdf"
        os.makedirs(os.path.dirname(out_path), exist_ok=True)

        c = canvas.Canvas(out_path, pagesize=letter)
        width, height = letter
        left_margin = 40
        top_margin = height - 40
        line_height = 14

        text_obj = c.beginText(left_margin, top_margin)
        text_obj.setFont("Times-Roman", 12)

        y = top_margin
        for para in doc.paragraphs:
            lines = para.text.splitlines() if para.text else [""]
            for line in lines:
                if y < 40:
                    c.drawText(text_obj)
                    c.showPage()
                    text_obj = c.beginText(left_margin, top_margin)
                    text_obj.setFont("Times-Roman", 12)
                    y = top_margin
                text_obj.textLine(line)
                y -= line_height

        c.drawText(text_obj)
        c.save()

        if not os.path.isfile(out_path):
            flash("Failed to create PDF file.", "error")
            return redirect(url_for("translator"))

        conversions_collection.insert_one({
            "type": "docx-to-pdf",
            "src_filename": in_name,
            "out_filename": os.path.basename(out_path),
            "created_at": datetime.now(timezone.utc),
            "owner_id": session.get("user_id")
        })

        return send_file(out_path, as_attachment=True, download_name=os.path.basename(out_path))
    except Exception as e:
        flash(f"DOCX to PDF failed: {e}", "error")
        return redirect(url_for("translator"))

@app.route("/image-to-pdf", methods=["POST"])
def image_to_pdf():
    image_file = request.files.get("image_file")
    if not image_file or image_file.filename == "":
        flash("Please select an image file.", "error")
        return redirect(url_for("translator"))

    in_name = secure_filename(image_file.filename)
    in_path = os.path.join(app.config["UPLOAD_FOLDER"], in_name)
    os.makedirs(os.path.dirname(in_path), exist_ok=True)
    image_file.save(in_path)

    try:
        with Image.open(in_path) as im:
            if im.mode in ("RGBA", "P"):
                im = im.convert("RGB")
            out_path = os.path.splitext(in_path)[0] + ".pdf"
            os.makedirs(os.path.dirname(out_path), exist_ok=True)
            im.save(out_path, "PDF", resolution=300)

        if not os.path.isfile(out_path):
            flash("Failed to create PDF file.", "error")
            return redirect(url_for("translator"))

        conversions_collection.insert_one({
            "type": "image-to-pdf",
            "src_filename": in_name,
            "out_filename": os.path.basename(out_path),
            "created_at": datetime.now(timezone.utc),
            "owner_id": session.get("user_id")
        })

        return send_file(out_path, as_attachment=True, download_name=os.path.basename(out_path))
    except Exception as e:
        flash(f"Image to PDF failed: {e}", "error")
        return redirect(url_for("translator"))

# ---------------- Main ----------------
if __name__ == "__main__":
    # Disable reloader on Windows to avoid MediaRecorder issues
    app.run(debug=True, use_reloader=False, host="0.0.0.0", port=int(os.getenv("PORT", "5000")))
